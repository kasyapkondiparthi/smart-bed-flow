from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Helper: set paragraph shading ────────────────────────────
def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)

# ── Helper: shade table cell ──────────────────────────────────
def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── Helper: add horizontal rule ───────────────────────────────
def add_hr(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1E40AF')
    pb.append(bottom)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

# ══════════════════════════════════════════════════════════════
# COVER BLOCK
# ══════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(cover, '1E3A8A')
run = cover.add_run('\n  SMART HOSPITAL BED ALLOCATION SYSTEM\n')
run.font.size  = Pt(22)
run.font.bold  = True
run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
run.font.name  = 'Calibri'
cover.paragraph_format.space_before = Pt(10)
cover.paragraph_format.space_after  = Pt(4)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(sub, '1E40AF')
sr = sub.add_run(f'  Project Work Summary  |  Generated: {datetime.datetime.now().strftime("%d %B %Y, %I:%M %p")}  ')
sr.font.size  = Pt(11)
sr.font.color.rgb = RGBColor(0xBF, 0xDB, 0xFF)
sr.font.name  = 'Calibri'
sub.paragraph_format.space_after = Pt(14)

# ══════════════════════════════════════════════════════════════
# SECTION helper
# ══════════════════════════════════════════════════════════════
def section_heading(doc, number, title):
    add_hr(doc)
    p = doc.add_paragraph()
    shade_paragraph(p, 'DBEAFE')
    r = p.add_run(f'  {number}. {title}  ')
    r.font.size  = Pt(13)
    r.font.bold  = True
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    r.font.name  = 'Calibri'
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)

def sub_heading(doc, title):
    p = doc.add_paragraph()
    r = p.add_run(f'▸  {title}')
    r.font.size  = Pt(11)
    r.font.bold  = True
    r.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    r.font.name  = 'Calibri'
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)

def body(doc, text):
    p = doc.add_paragraph(style='Normal')
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.name = 'Calibri'
    p.paragraph_format.left_indent   = Cm(0.5)
    p.paragraph_format.space_after   = Pt(2)

def bullet(doc, text):
    p = doc.add_paragraph(style='Normal')
    r = p.add_run(f'  •  {text}')
    r.font.size = Pt(10.5)
    r.font.name = 'Calibri'
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_after  = Pt(1)

def table3(doc, headers, rows, col_colors=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header row
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        shade_cell(hdr[i], '1E40AF')
        run = hdr[i].paragraphs[0].runs[0]
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size  = Pt(10)
        run.font.name  = 'Calibri'
    # data rows
    for ri, row in enumerate(rows):
        cells = t.rows[ri+1].cells
        bg = 'F0F7FF' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cells[ci].text = val
            shade_cell(cells[ci], bg)
            run = cells[ci].paragraphs[0].runs[0]
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ══════════════════════════════════════════════════════════════
# 1. UI/UX OVERHAUL
# ══════════════════════════════════════════════════════════════
section_heading(doc, 1, 'UI / UX Overhaul')

sub_heading(doc, 'Global Theme System')
for b in [
    'Implemented Light / Dark / System theme modes across ALL pages.',
    'Created src/utils/theme.ts with applyTheme() and loadTheme() utilities.',
    'Applied theme at root HTML level via document.documentElement.classList.',
    'Flicker-free theme loading on app start via main.tsx.',
]:
    bullet(doc, b)

sub_heading(doc, 'Landing Page (Welcome.tsx)')
for b in [
    'Premium glassmorphic hero design with animated CSS grid overlay.',
    'Dynamic gradient background: Light = soft blue-white, Dark = deep navy.',
    'Theme switcher (Sun / Moon / Monitor icons) embedded in the header.',
    'Removed all emoji / SVG section icons for a clean, text-first layout.',
]:
    bullet(doc, b)

sub_heading(doc, 'Login Page (Login.tsx)')
for b in [
    'Full redesign into a "Hospital OS" mission-critical entry point.',
    'Glassmorphism panels with radial gradient backgrounds.',
    'Dedicated theme switcher inside the login console.',
]:
    bullet(doc, b)

sub_heading(doc, 'Dashboard Pages')
for b in [
    'Command Center, Patient Allocation, Shift Audit Logs — all fully theme-aware.',
    'Card-based layouts with proper light/dark contrast ratios.',
    'Consistent spacing, shadows, and Calibri/Inter typography.',
]:
    bullet(doc, b)

# ══════════════════════════════════════════════════════════════
# 2. INTELLIGENT PATIENT BOOKING ENGINE
# ══════════════════════════════════════════════════════════════
section_heading(doc, 2, 'Intelligent Patient Booking Engine')

sub_heading(doc, 'Self-Booking Portal Features')
table3(doc,
    ['Feature', 'Detail'],
    [
        ['5-Second Processing Delay', 'Simulates real-world medical data synchronization'],
        ['Loading State',             'Button shows "⏳ Processing booking..." during delay'],
        ['Duplicate Detection',       'Blocks re-registration of the same patient name via DB query'],
        ['Priority Split',            'Critical = Instant ICU; Others = Pending Triage Review'],
        ['Clinical Severity Levels',  '🔴 Critical Emergency / 🟡 Moderate Trace / 🟢 Low Priority'],
        ['ICU Toggle',                'Override switch to request ICU admission directly'],
    ]
)

sub_heading(doc, 'Critical Patient Flow 🔴')
for b in [
    'User selects "Critical Emergency" severity.',
    'No wait time — instant dispatch bypasses all queue logic.',
    'ICU bed auto-assigned with short ID (e.g., ICU-A3B2).',
    'Patient Portal shows confirmation card: Bed No. / Floor / Room.',
    'DB record inserted as status: "Confirmed", assigned_bed: "ICU".',
]:
    bullet(doc, b)

sub_heading(doc, 'Non-Critical Patient Flow 🟡🟢')
for b in [
    'User submits form → 5-second processing delay begins.',
    'Duplicate check runs against the patients table.',
    'Patient inserted into DB as status: "Waiting", assigned_bed: "Waiting".',
    'Patient Portal shows "Pending Triage Review" card with Batch Allocate note.',
    'Staff sees the patient in Patient Allocation → Pending Triage Review panel.',
    'Staff runs Batch Allocate → system assigns bed and updates status to Confirmed.',
]:
    bullet(doc, b)

# ══════════════════════════════════════════════════════════════
# 3. PATIENT ALLOCATION PAGE
# ══════════════════════════════════════════════════════════════
section_heading(doc, 3, 'Patient Allocation Page Upgrades')

sub_heading(doc, 'Pending Triage Review Panel')
table3(doc,
    ['Item', 'Before', 'After'],
    [
        ['Filter',        'Only status = "Pending Approval"',   'status = "Waiting" OR "Pending Approval"'],
        ['Sort Order',    'Chronological only',                  'Critical first, then by arrival time'],
        ['Patient Cards', 'Plain name + severity text',          'Color-coded severity badges + ICU tag + Self-Booked label'],
        ['Counter Badge', '"X NEW"',                            '"X PENDING" (live count)'],
        ['Actions',       'Approve / Reject buttons',            'Same + tooltips + rounded styling'],
    ]
)

sub_heading(doc, 'Patient Registry Table')
table3(doc,
    ['Improvement', 'Detail'],
    [
        ['Action Buttons',   'Always visible (previously hidden until hover)'],
        ['Icons',            'Proper Lucide Edit2 / Trash2 icons (was emoji ✏️ 🗑)'],
        ['Status Badge',     '"Waiting" now displays as "Pending Triage" with pulsing amber color'],
        ['Bed Column',       'Now shows: Bed Number + Floor + Room Number stacked'],
    ]
)

# ══════════════════════════════════════════════════════════════
# 4. SHIFT AUDIT LOGS — CALENDAR FIX
# ══════════════════════════════════════════════════════════════
section_heading(doc, 4, 'Shift Audit Logs — Calendar Date Fix')

sub_heading(doc, 'Root Cause')
body(doc, 'The old code stored the selected date as a string ("2026-03-31") and parsed it '
     'with new Date("2026-03-31"), which resolves to UTC midnight. In India (UTC+5:30), this '
     'meant selecting March 31 actually filtered for March 30, returning zero results.')

sub_heading(doc, 'Fix Applied')
table3(doc,
    ['Area', 'Before', 'After'],
    [
        ['State type',         'string ("2026-03-31")',           'Date object (timezone-safe)'],
        ['Date comparison',    'isSameDay(new Date(string)) — UTC bug', 'toLocaleDateString() — always local'],
        ['Calendar selected',  'parseISO(string)',                'Direct Date object (no conversion)'],
        ['Popover close',      'Always closes on any click',      'Only closes when a valid date is picked'],
        ['Bonus feature',      '—',                               '"Today" quick-select button added'],
    ]
)

# ══════════════════════════════════════════════════════════════
# 5. DATABASE & SUPABASE FIXES
# ══════════════════════════════════════════════════════════════
section_heading(doc, 5, 'Database & Supabase Fixes')

table3(doc,
    ['Error', 'Root Cause', 'Fix Applied'],
    [
        ['PGRST204',           'Schema cache stale after ALTER TABLE',           'NOTIFY pgrst, "reload schema"'],
        ['varchar(10) overflow','Bed IDs like ICU-TEMP-1743435922000 = 22 chars', 'Short IDs: ICU-A3B2 (8 chars)'],
        ['status col not found','PostgREST cache out of sync',                    'Schema reload + ALTER COLUMN TYPE varchar(20)'],
        ['RLS blocking inserts', 'No public insert policy defined',               'CREATE POLICY for SELECT/INSERT/UPDATE/DELETE'],
        ['Registry Insert Failed','Recently added columns not in cache',          'Simplified insert to stable columns only'],
    ]
)

sub_heading(doc, 'Safe Insert Payload (Final Version)')
body(doc, 'The booking insert now only sends columns that PostgREST reliably knows:')
for b in ['name (text)', 'severity ("Low" | "Moderate" | "Critical")',
          'needs_icu (boolean)', 'assigned_bed ("Waiting" | "ICU" | "Normal")']:
    bullet(doc, b)

# ══════════════════════════════════════════════════════════════
# 6. NEURAL SYNC BACKGROUND ENGINE
# ══════════════════════════════════════════════════════════════
section_heading(doc, 6, 'Neural Sync Background Engine')
for b in [
    '20-second auto-sync: Background setInterval calls fetchPatientsList(true) every 20 seconds.',
    'Triggers runGlobalReallocation() in PatientContext automatically.',
    'Promotes "Waiting" patients to "Confirmed" as beds free up — zero staff intervention required.',
    'Removed duplicate redundant useEffect blocks that were calling usePatients() inside an effect (invalid React hook usage).',
]:
    bullet(doc, b)

# ══════════════════════════════════════════════════════════════
# 7. KEY FILES MODIFIED
# ══════════════════════════════════════════════════════════════
section_heading(doc, 7, 'Key Files Modified')

table3(doc,
    ['File', 'What Changed'],
    [
        ['src/pages/Welcome.tsx',                   'Entire booking engine, dual-track triage, Neural Sync, clinical severity levels'],
        ['src/pages/Login.tsx',                     'Full premium Hospital OS redesign with glassmorphism'],
        ['src/pages/dashboard/AllocationPage.tsx',  'Pending Triage Review panel — shows Waiting + Pending Approval'],
        ['src/pages/dashboard/HistoryPage.tsx',     'Calendar timezone fix — Date object state, toLocaleDateString()'],
        ['src/components/AllocationTable.tsx',      'Always-visible action buttons, room number in Bed column, status badge'],
        ['src/context/PatientContext.tsx',           'Global reallocation engine, bed mapping normalization'],
        ['src/utils/theme.ts',                      'NEW FILE — Theme utility (applyTheme, loadTheme, getTheme)'],
        ['server/index.js',                         'NEW FILE — Backend proxy server for AI chatbot integration'],
        ['src/App.tsx',                             'JSX indentation cleanup, route structure'],
        ['src/components/HospitalHeader.tsx',       'Theme-aware header with switcher'],
        ['src/components/Sidebar.tsx',              'Dark/light mode sidebar styling'],
        ['src/index.css',                           'Global CSS tokens, glassmorphism utilities'],
    ]
)

# ══════════════════════════════════════════════════════════════
# 8. END-TO-END SYSTEM FLOW
# ══════════════════════════════════════════════════════════════
section_heading(doc, 8, 'End-to-End System Flow')

flow = [
    ('Patient visits Landing Page', 'Opens Patient Portal via "Enter Patient Portal" button'),
    ('Fills booking form',          'Name, Clinical Severity, ICU toggle'),
    ('Clicks "Book Medical Space"', '5-second processing delay begins'),
    ('Duplicate check',             'Database query to prevent double registration'),
    ('Critical branch',             'Instant ICU allocation → Confirmed card with Bed/Floor/Room'),
    ('Non-critical branch',         'status: Waiting inserted → Pending Triage Review card'),
    ('Staff opens Allocation page', 'Pending Triage Review panel shows all waiting patients'),
    ('Staff clicks Batch Allocate', 'runGlobalReallocation() assigns beds by priority'),
    ('Patient confirmed',           'status → Confirmed, bed assigned in Patient Registry'),
    ('Events logged',               'All actions appear in Shift Audit Logs with date-filter support'),
]
table3(doc, ['Step', 'Action'], [[str(i+1)+'. '+s, d] for i,(s,d) in enumerate(flow)])

# ══════════════════════════════════════════════════════════════
# 9. GIT COMMIT SUMMARY
# ══════════════════════════════════════════════════════════════
section_heading(doc, 9, 'Git Repository — Final Push')
table3(doc,
    ['Field', 'Value'],
    [
        ['Repository', 'https://github.com/kasyapkondiparthi/smart-bed-flow.git'],
        ['Branch',     'main'],
        ['Commit',     'feat: Intelligent Patient Booking System & UI Enhancements'],
        ['Files Changed', '26 files'],
        ['Insertions', '3,232 lines added'],
        ['Deletions',  '923 lines removed'],
        ['New Files',  'server/index.js, server/package.json, src/utils/theme.ts'],
    ]
)

# ── Footer ─────────────────────────────────────────────────────
add_hr(doc)
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run(f'Smart Hospital Bed Allocation System  •  Project Summary  •  {datetime.datetime.now().strftime("%d %B %Y")}')
fr.font.size  = Pt(9)
fr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
fr.font.name  = 'Calibri'

# ── Save ───────────────────────────────────────────────────────
out = '/Users/venkatasrikasyapkondiparthi/Downloads/SmartBedFlow_ProjectSummary.docx'
doc.save(out)
print(f'✅  Saved → {out}')
